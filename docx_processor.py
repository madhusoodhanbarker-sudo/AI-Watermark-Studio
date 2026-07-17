from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_text_watermark(
    input_docx,
    output_docx,
    text="CONFIDENTIAL",
    font_size=28
):
    # Open Word document
    doc = Document(input_docx)

    # Add watermark to each section header
    for section in doc.sections:
        header = section.header

        # Use existing paragraph or create one
        if header.paragraphs:
            paragraph = header.paragraphs[0]
        else:
            paragraph = header.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = True

    # Save document
    doc.save(output_docx)

    return output_docx
